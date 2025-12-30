from docx import Document

def generate_docx(sections: dict, output_path: str):
    doc = Document()

    for title, content in sections.items():
        doc.add_heading(title.upper(), level=2)
        doc.add_paragraph(content)

    doc.save(output_path)
